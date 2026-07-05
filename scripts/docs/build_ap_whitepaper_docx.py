from __future__ import annotations

import math
import random
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


AP_ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = AP_ROOT / "docs"
SOURCE_TXT = DOCS_DIR / "AP_Bottom_Principles_Whitepaper_20260626.txt"
OUTPUT_DOCX = DOCS_DIR / "AP底层原理白皮书_v0.4_最终理论纠偏精排版_20260626.docx"
FIGURE_DIR = DOCS_DIR / "assets" / "ap_whitepaper_figures"


PALETTE = {
    "ink": "1F2937",
    "muted": "6B7280",
    "line": "CBD5E1",
    "blue": "2E4960",
    "blue2": "4E7FA3",
    "gold": "E8C547",
    "green": "4F8A5B",
    "red": "C15A4A",
    "orange": "D58A35",
    "purple": "7566A0",
    "cyan": "4C9A9A",
    "paper": "F7F7F8",
    "soft": "EFF3F6",
    "white": "FFFFFF",
}


VOLUME_RE = re.compile(r"^第[一二三四五六七八九十百〇零0-9]+卷\b")
H2_RE = re.compile(r"^\d{1,3}\.\s+\S")
H3_RE = re.compile(r"^\d{1,3}\.\d+[a-zA-Z]?\s+\S")
CODE_HINT_RE = re.compile(r"[={}<>^λ\\]|->|=>|clamp|lambda|Recall|Propagate|StructurePool|Field_|C\*|Cstar")


def rgb(hex_color: str) -> tuple[int, int, int]:
    hex_color = hex_color.strip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def find_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/Dengb.ttf") if bold else Path("C:/Windows/Fonts/Deng.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_TITLE = find_font(46, bold=True)
FONT_H = find_font(28, bold=True)
FONT_M = find_font(22)
FONT_S = find_font(18)
FONT_XS = find_font(15)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        trial = current + char
        box = draw.textbbox((0, 0), trial, font=font)
        if box[2] - box[0] <= width or not current:
            current = trial
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_text_center(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str = PALETTE["ink"],
    line_gap: int = 4,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 24)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) // 2
    for line, h in zip(lines, heights):
        text_box = draw.textbbox((0, 0), line, font=font)
        w = text_box[2] - text_box[0]
        draw.text((x1 + (x2 - x1 - w) // 2, y), line, font=font, fill=rgb(fill))
        y += h + line_gap


def draw_text_left(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str = PALETTE["ink"],
    width: int = 500,
    line_gap: int = 6,
) -> int:
    x, y = xy
    for line in wrap_text(draw, text, font, width):
        draw.text((x, y), line, font=font, fill=rgb(fill))
        y += draw.textbbox((0, 0), line, font=font)[3] + line_gap
    return y


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    label: str,
    fill: str,
    outline: str = PALETTE["blue"],
    font: ImageFont.FreeTypeFont = FONT_M,
    radius: int = 24,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=rgb(fill), outline=rgb(outline), width=3)
    draw_text_center(draw, box, label, font)


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = PALETTE["blue"],
    width: int = 5,
    head: int = 18,
) -> None:
    sx, sy = start
    ex, ey = end
    draw.line((sx, sy, ex, ey), fill=rgb(color), width=width)
    angle = math.atan2(ey - sy, ex - sx)
    left = (ex - head * math.cos(angle - math.pi / 6), ey - head * math.sin(angle - math.pi / 6))
    right = (ex - head * math.cos(angle + math.pi / 6), ey - head * math.sin(angle + math.pi / 6))
    draw.polygon([(ex, ey), left, right], fill=rgb(color))


def new_canvas(title: str, subtitle: str | None = None) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1800, 1050), rgb(PALETTE["white"]))
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 120), fill=rgb(PALETTE["blue"]))
    draw.text((70, 32), title, font=FONT_TITLE, fill=rgb(PALETTE["white"]))
    if subtitle:
        draw.text((70, 85), subtitle, font=FONT_S, fill=rgb("DDE9F1"))
    draw.rectangle((0, 1028, 1800, 1050), fill=rgb(PALETTE["gold"]))
    return img, draw


def save_figure(img: Image.Image, name: str) -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / name
    img.save(path, "PNG", optimize=True)
    return path


def figure_full_loop() -> Path:
    img, draw = new_canvas("图 1  AP 完整认知闭环", "从感受器到行动反馈，每个 tick 都回到同一条生命流")
    boxes = [
        ((90, 210, 350, 335), "外界/内源输入\n文本 视觉 听觉\n行动反馈 奖惩", "EFF6FF"),
        ((430, 210, 690, 335), "感受器\n多模态 occurrence", "F0FDF4"),
        ((770, 210, 1030, 335), "状态池\n无序认知场", "FFF7ED"),
        ((1110, 210, 1370, 335), "短期结构池\n序列/空间/图/时空流", "F5F3FF"),
        ((1450, 210, 1710, 335), "统一经验流\n追加事件与索引", "F8FAFC"),
        ((1110, 470, 1370, 595), "B / C / C*\n现状召回 预测 追溯", "ECFEFF"),
        ((770, 470, 1030, 595), "认知感受\n惊 违和 合理 把握", "FEFCE8"),
        ((430, 470, 690, 595), "情绪慢量\n疲劳 期待 压力", "FDF2F8"),
        ((90, 470, 350, 595), "行动竞争\n看 写 改 提交 停", "F0FDFA"),
        ((430, 725, 1370, 850), "行动结果再次成为输入：DraftGrid、视焦点移动、教师反馈、工具执行都回到下一 tick", "F7F7F8"),
    ]
    for box, label, fill in boxes:
        draw_box(draw, box, label, fill)
    arrows = [
        ((350, 272), (430, 272)),
        ((690, 272), (770, 272)),
        ((1030, 272), (1110, 272)),
        ((1370, 272), (1450, 272)),
        ((1240, 335), (1240, 470)),
        ((1110, 532), (1030, 532)),
        ((770, 532), (690, 532)),
        ((430, 532), (350, 532)),
        ((220, 595), (600, 725)),
        ((1280, 725), (1590, 335)),
        ((600, 725), (220, 335)),
    ]
    for start, end in arrows:
        draw_arrow(draw, start, end)
    draw_text_left(
        draw,
        (95, 905),
        "核心读法：AP 不是先识别再回答的流水线，而是在每个 tick 内让“现实、预测、回忆、行动、奖惩”相互调制。白箱价值来自每个环节都能回放。",
        FONT_M,
        PALETTE["blue"],
        1600,
    )
    return save_figure(img, "fig01_ap_full_loop.png")


def figure_statepool_ssp() -> Path:
    img, draw = new_canvas("图 2  状态池与短期结构池", "StatePool 是类型投影，SSP 是 occurrence 流；重复、计数和顺序来自流")
    draw_box(draw, (95, 210, 760, 820), "状态池 StatePool\n无序场：相同 SA 类型会聚合\n保存当前实能量 R、虚能量 V、注意 A、疲劳 F\n适合表示“此刻有哪些东西在心里亮着”", "EFF6FF", PALETTE["blue"], FONT_M)
    draw_box(draw, (1040, 210, 1705, 820), "短期结构池 SSP\n有序/空间/图/时空 occurrence 流\n同一对象在不同时间或空间位置不合并\n适合表示“刚才发生了几次、顺序是什么、在哪里”", "F0FDF4", PALETTE["green"], FONT_M)
    draw_arrow(draw, (760, 505), (1040, 505), PALETTE["orange"], width=7)
    draw_text_center(draw, (790, 430, 1010, 480), "注意调制后写入", FONT_S, PALETTE["orange"])
    draw_arrow(draw, (1040, 600), (760, 600), PALETTE["purple"], width=5)
    draw_text_center(draw, (790, 625, 1010, 675), "召回/预测回灌", FONT_S, PALETTE["purple"])
    for i, x in enumerate(range(1120, 1580, 95)):
        fill = ["FFF7ED", "F5F3FF", "ECFEFF", "FEFCE8", "FDF2F8"][i % 5]
        draw.rounded_rectangle((x, 655, x + 70, 725), radius=16, fill=rgb(fill), outline=rgb(PALETTE["line"]), width=2)
        draw_text_center(draw, (x, 655, x + 70, 725), ["你", "好", "啊", "图像", "奖励"][i % 5], FONT_XS)
        if i < 4:
            draw_arrow(draw, (x + 70, 690), (x + 95, 690), PALETTE["muted"], width=3, head=10)
    for i, (label, y) in enumerate([("苹果", 335), ("问候", 430), ("未闭合", 525), ("视焦点", 620)]):
        draw.ellipse((260, y, 360, y + 100), fill=rgb(["FFF7ED", "F0FDF4", "FDF2F8", "F5F3FF"][i]), outline=rgb(PALETTE["line"]), width=2)
        draw_text_center(draw, (235, y, 385, y + 100), label, FONT_XS)
    draw_text_left(
        draw,
        (110, 890),
        "这张图对应最新裁定：重复感、累计次数、列竖式、数手指、视觉空间关系，都不靠状态池里的“合并字段”硬做，而靠 SSP 中 occurrence 的结构关系自然得到。",
        FONT_M,
        PALETTE["blue"],
        1580,
    )
    return save_figure(img, "fig02_statepool_ssp.png")


def figure_recall_cstar() -> Path:
    img, draw = new_canvas("图 3  B / C / C* 召回与理解", "当前结构召回历史结构波，再向后预测、向前追溯，合成 C* 回灌")
    draw_box(draw, (90, 280, 360, 430), "当前 SSP 查询\nQ_t", "EFF6FF")
    draw_box(draw, (520, 185, 870, 315), "历史相似结构 1\nB_t: 现状认知波", "F8FAFC")
    draw_box(draw, (520, 370, 870, 500), "历史相似结构 2\nB_t: 现状认知波", "F8FAFC")
    draw_box(draw, (1040, 185, 1390, 315), "向后传播\nC_forward\n预测接下来会怎样", "ECFEFF", PALETTE["cyan"])
    draw_box(draw, (1040, 370, 1390, 500), "向前传播\nC_backward\n解释为什么会这样", "FFF7ED", PALETTE["orange"])
    draw_box(draw, (1500, 280, 1710, 430), "C* 包\n预测槽\n解释槽\n反例槽", "F5F3FF", PALETTE["purple"])
    draw_arrow(draw, (360, 355), (520, 250))
    draw_arrow(draw, (360, 355), (520, 435))
    draw_arrow(draw, (870, 250), (1040, 250), PALETTE["cyan"])
    draw_arrow(draw, (870, 435), (1040, 435), PALETTE["orange"])
    draw_arrow(draw, (1390, 250), (1500, 335), PALETTE["purple"])
    draw_arrow(draw, (1390, 435), (1500, 380), PALETTE["purple"])
    draw_box(draw, (620, 680, 1570, 805), "C* 回灌状态池：形成惊讶、合理感、把握感、期待、压力、闭合/未闭合，并影响下一 tick 的注意与行动", "FEFCE8", PALETTE["gold"])
    draw_arrow(draw, (1605, 430), (1200, 680), PALETTE["gold"], width=6)
    draw_text_left(
        draw,
        (100, 895),
        "拟人读法：人会先觉得“像以前某件事”，再预测后续，也会反向找原因。伪因果不是被天生禁止，而是在后天反例、失败行动和教师纠正中逐步松动。",
        FONT_M,
        PALETTE["blue"],
        1600,
    )
    return save_figure(img, "fig03_recall_cstar.png")


def figure_foveated_vision() -> Path:
    img, draw = new_canvas("图 4  视焦点采样与内心画面重建", "焦点附近高清，周边低概率稀疏；多 tick 逐步补全 sensory canvas")
    random.seed(7)
    left = (110, 210, 830, 820)
    right = (970, 210, 1690, 820)
    draw.rounded_rectangle(left, radius=24, fill=rgb("F8FAFC"), outline=rgb(PALETTE["line"]), width=3)
    draw.rounded_rectangle(right, radius=24, fill=rgb("F8FAFC"), outline=rgb(PALETTE["line"]), width=3)
    draw_text_center(draw, (120, 225, 820, 275), "视觉输入上的主动采样", FONT_M, PALETTE["blue"])
    draw_text_center(draw, (980, 225, 1680, 275), "状态池视觉 SA 重建的内心画面", FONT_M, PALETTE["blue"])
    focus = (470, 515)
    for _ in range(900):
        x = random.randint(left[0] + 45, left[2] - 45)
        y = random.randint(left[1] + 90, left[3] - 45)
        d = math.hypot(x - focus[0], y - focus[1])
        p = max(0.08, 1.0 - d / 520)
        if random.random() < p:
            r = 2 if d > 240 else 4 if d > 110 else 7
            col = random.choice(["C15A4A", "D58A35", "E8C547", "4F8A5B", "2E4960"])
            draw.ellipse((x - r, y - r, x + r, y + r), fill=rgb(col))
    draw.ellipse((focus[0] - 95, focus[1] - 95, focus[0] + 95, focus[1] + 95), outline=rgb(PALETTE["red"]), width=8)
    draw.ellipse((focus[0] - 20, focus[1] - 20, focus[0] + 20, focus[1] + 20), outline=rgb(PALETTE["red"]), width=5)
    draw_text_center(draw, (360, 620, 590, 675), "视焦点\nclarity field", FONT_S, PALETTE["red"])
    focus2 = (1320, 505)
    for _ in range(650):
        x = random.randint(right[0] + 55, right[2] - 55)
        y = random.randint(right[1] + 95, right[3] - 55)
        d = math.hypot(x - focus2[0], y - focus2[1])
        alpha_like = max(0.18, 1.0 - d / 520)
        if random.random() < alpha_like:
            base_r = 2 if d > 250 else 5 if d > 120 else 9
            col = random.choice(["D97A62", "E8C547", "6AA76D", "6B8FA8", "A76A6A"])
            draw.ellipse((x - base_r, y - base_r, x + base_r, y + base_r), fill=rgb(col))
    for radius, color in [(180, "CBD5E1"), (95, "4E7FA3"), (28, "2E4960")]:
        draw.ellipse((focus2[0] - radius, focus2[1] - radius, focus2[0] + radius, focus2[1] + radius), outline=rgb(color), width=4)
    draw_text_left(draw, (1040, 695), "R_sketch: 现实采样\nV_sketch: 预测补全\nsource mask: 来源可追溯", FONT_S, PALETTE["ink"], 520)
    draw_text_left(
        draw,
        (105, 895),
        "实现要点：内心画面不从原图直接贴图，也不只画一个椭圆；它来自每 tick 进入状态池的视觉 patch payload、clarity、R/V 能量、source mask 与多 tick 累积。",
        FONT_M,
        PALETTE["blue"],
        1600,
    )
    return save_figure(img, "fig04_foveated_vision.png")


def figure_draftgrid() -> Path:
    img, draw = new_canvas("图 5  DraftGrid 与二维草稿", "输出不是一次生成整句，而是行动器在草稿栏中逐 tick 写、看、改、提交")
    grid_left, grid_top = 190, 250
    cell_w, cell_h = 110, 70
    rows, cols = 5, 9
    draw.rounded_rectangle((120, 190, 1170, 695), radius=26, fill=rgb("F8FAFC"), outline=rgb(PALETTE["line"]), width=3)
    for r in range(rows):
        for c in range(cols):
            x1 = grid_left + c * cell_w
            y1 = grid_top + r * cell_h
            draw.rectangle((x1, y1, x1 + cell_w, y1 + cell_h), outline=rgb("D5DEE8"), width=2)
    tokens = {(0, 0): "嗯", (0, 1): ",", (0, 2): "是", (0, 3): "苹", (0, 4): "果", (1, 0): "可", (1, 1): "以", (1, 2): "再", (1, 3): "看", (2, 2): "改"}
    for (r, c), token in tokens.items():
        x1 = grid_left + c * cell_w
        y1 = grid_top + r * cell_h
        draw_text_center(draw, (x1, y1, x1 + cell_w, y1 + cell_h), token, FONT_H, PALETTE["blue"])
    draw.rounded_rectangle((grid_left + 3 * cell_w, grid_top - 12, grid_left + 5 * cell_w, grid_top + cell_h + 12), radius=10, outline=rgb(PALETTE["red"]), width=5)
    actions = [
        ((1260, 220, 1650, 310), "write_cell\n写入候选片段", "EFF6FF"),
        ((1260, 350, 1650, 440), "review_grid\n回看自己写了什么", "F0FDF4"),
        ((1260, 480, 1650, 570), "revise_cell\n修改/删除/换位置", "FFF7ED"),
        ((1260, 610, 1650, 700), "commit_reply\n提交成为外部输出", "F5F3FF"),
    ]
    for box, label, fill in actions:
        draw_box(draw, box, label, fill, font=FONT_S)
    for y in [265, 395, 525, 655]:
        draw_arrow(draw, (1170, y), (1260, y), PALETTE["blue"], width=4, head=14)
    draw_box(draw, (470, 770, 1430, 880), "提交前后的草稿内容会作为 SELF_DRAFT_GRID 感受器输入回到 SSP，使 AP 能“看见自己刚写了什么”，再决定继续、修改、停下或发送。", "FEFCE8", PALETTE["gold"], FONT_M)
    draw_arrow(draw, (960, 695), (960, 770), PALETTE["gold"], width=6)
    return save_figure(img, "fig05_draftgrid.png")


def figure_experience_log() -> Path:
    img, draw = new_canvas("图 6  统一经验流与加速索引", "真相源是 append-only 经验流；Zvec/ANN/倒排表/记忆包是可重建加速与分享层")
    y = 320
    xs = [130, 360, 590, 820, 1050, 1280]
    labels = ["tick 1\n事件增量", "tick 2\n感受/行动", "tick 3\n奖励/惩罚", "tick 4\n教师纠正", "tick 5\n预测验证", "tick 6\n焦点片段"]
    for x, label in zip(xs, labels):
        draw_box(draw, (x, y, x + 165, y + 115), label, "F8FAFC", font=FONT_S, radius=18)
    for i in range(len(xs) - 1):
        draw_arrow(draw, (xs[i] + 165, y + 58), (xs[i + 1], y + 58), PALETTE["blue"], width=4, head=14)
    draw_box(draw, (190, 580, 550, 720), "派生索引\nANN / Zvec / rolling hash\n只加速召回", "EFF6FF", PALETTE["blue"], FONT_S)
    draw_box(draw, (720, 580, 1080, 720), "记忆包\n经验子图 + 元数据\n可导入/导出/卸载", "F0FDF4", PALETTE["green"], FONT_S)
    draw_box(draw, (1250, 580, 1610, 720), "审计旁路\n近期快照/图表/回放\n不替代核心真相", "FFF7ED", PALETTE["orange"], FONT_S)
    draw_arrow(draw, (480, 435), (370, 580), PALETTE["blue"], width=4, head=14)
    draw_arrow(draw, (900, 435), (900, 580), PALETTE["green"], width=4, head=14)
    draw_arrow(draw, (1280, 435), (1430, 580), PALETTE["orange"], width=4, head=14)
    draw_text_left(
        draw,
        (120, 850),
        "积极证明口径：只要索引和记忆包可以由经验流重建，工程加速就不会改变 AP-Core 的语义来源；回放、分享和性能优化可以共存。",
        FONT_M,
        PALETTE["blue"],
        1580,
    )
    return save_figure(img, "fig06_experience_log.png")


def figure_roadmap() -> Path:
    img, draw = new_canvas("图 7  Phase20.7 实施路线", "先真闭环，再视觉重建，再工作台；页面只显示真实 RuntimeTickEvent")
    stages = [
        ("Stage 0", "红线隔离\n清旧投影/命中/整图路径"),
        ("Stage 1", "StatePool + SSP\n最小 tick 闭环"),
        ("Stage 2", "统一经验流\n追加事件与索引"),
        ("Stage 3", "B/C/C*\n预测与追溯"),
        ("Stage 4", "DraftGrid\n逐 tick 写看改发"),
        ("Stage 5", "视觉/听觉重建\npatch payload + focus"),
        ("Stage 6", "工作台 UI\n只读真实 trace"),
    ]
    x = 85
    for idx, (name, desc) in enumerate(stages):
        fill = ["EFF6FF", "F0FDF4", "FFF7ED", "F5F3FF", "ECFEFF", "FEFCE8", "FDF2F8"][idx]
        draw.rounded_rectangle((x, 300, x + 220, 590), radius=24, fill=rgb(fill), outline=rgb(PALETTE["blue"]), width=3)
        draw_text_center(draw, (x, 325, x + 220, 390), name, FONT_H, PALETTE["blue"])
        draw_text_center(draw, (x + 15, 415, x + 205, 565), desc, FONT_S, PALETTE["ink"])
        if idx < len(stages) - 1:
            draw_arrow(draw, (x + 220, 445), (x + 270, 445), PALETTE["blue"], width=4, head=14)
        x += 245
    draw_box(draw, (255, 735, 1545, 860), "验收原则：每一步都要能说明“可以证明什么”，并给出 RuntimeTickEvent、经验流增量、召回来源、行动竞争与回放证据。", "F8FAFC", PALETTE["gold"], FONT_M)
    return save_figure(img, "fig07_phase20_7_roadmap.png")


def build_figures() -> list[tuple[str, Path]]:
    return [
        ("AP 完整认知闭环", figure_full_loop()),
        ("状态池与短期结构池", figure_statepool_ssp()),
        ("B/C/C* 召回与理解", figure_recall_cstar()),
        ("视焦点采样与内心画面重建", figure_foveated_vision()),
        ("DraftGrid 与二维草稿", figure_draftgrid()),
        ("统一经验流与加速索引", figure_experience_log()),
        ("Phase20.7 实施路线", figure_roadmap()),
    ]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D9E2EC") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, east_asia: str = "Microsoft YaHei", ascii_font: str | None = None) -> None:
    ascii_font = ascii_font or east_asia
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 4, line: float = 1.18) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def request_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("第 ")
    add_field(paragraph, "PAGE")
    paragraph.add_run(" 页")
    for run in paragraph.runs:
        set_run_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(107, 114, 128)


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.9)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color, before, after in [
        ("Title", 28, PALETTE["blue"], 0, 12),
        ("Subtitle", 13, PALETTE["muted"], 0, 6),
        ("Heading 1", 19, PALETTE["blue"], 16, 8),
        ("Heading 2", 14, PALETTE["blue2"], 12, 6),
        ("Heading 3", 11.5, PALETTE["ink"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*rgb(color))
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = styles.add_style("AP Code Block", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8.6)
    code.font.color.rgb = RGBColor(31, 41, 55)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code.paragraph_format.left_indent = Cm(0.45)
    code.paragraph_format.right_indent = Cm(0.2)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.04

    caption = styles.add_style("AP Figure Caption", 1)
    caption.font.name = "Microsoft YaHei"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(107, 114, 128)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)

    return doc


def add_cover(doc: Document, version_line: str) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AP底层原理白皮书")
    set_run_font(run, "Microsoft YaHei")
    run.font.size = Pt(31)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*rgb(PALETTE["blue"]))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Artificial PsyArch Bottom Principles Whitepaper")
    set_run_font(run, "Microsoft YaHei")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*rgb(PALETTE["muted"]))

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rows = [
        ("版本", version_line.replace("版本:", "").strip()),
        ("定位", "AP底层哲学、数学流程、实现红线、证据边界的当前理论纠偏正本"),
        ("适用", "APV3.0test 数学模型重建 / APV2.1 证据解释 / GL 教学协议 / SNS 桌宠产品壳 / Agent 与未来具身路线"),
        ("读者", "小白读者、技术实现者、审查者、论文写作者"),
        ("生成", "由源 TXT 自动精排版生成，正文保持 30 卷完整结构"),
    ]
    for row, (key, value) in zip(table.rows, rows):
        row.cells[0].text = key
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "EAF1F6")
        set_cell_shading(row.cells[1], "F7F7F8")
        for cell in row.cells:
            set_cell_border(cell)
            for para in cell.paragraphs:
                set_paragraph_spacing(para, after=2)
                for run in para.runs:
                    set_run_font(run)
                    run.font.size = Pt(9.5)
            row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("核心口径：AP 的证明力来自同一条 tick 闭环，而不是前端投影、答案表或独立捷径。")
    set_run_font(run)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*rgb(PALETTE["blue"]))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("这份 Word 版增加导读、图示、公式块和清晰标题层级，方便作为论文底层纠偏与实现设计文档使用。")
    set_run_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(*rgb(PALETTE["muted"]))
    doc.add_page_break()


def add_callout(doc: Document, title: str, body: str, fill: str = "F7F7F8", accent: str = "2E4960") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.allow_autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_run_font(run)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*rgb(accent))
    run.font.size = Pt(11)
    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=2)
    for part in body.split("\n"):
        if part:
            r = p.add_run(part)
            set_run_font(r)
            r.font.size = Pt(9.5)
            p.add_run("\n")
    doc.add_paragraph()


def add_reader_guide(doc: Document) -> None:
    doc.add_heading("读者导读", level=1)
    p = doc.add_paragraph(
        "这份白皮书可以按两条路径阅读：先用图和导读建立直觉，再进入 30 卷正文查看定义、公式、示例、红线和实施路线。"
    )
    set_paragraph_spacing(p)

    add_callout(
        doc,
        "给小白的读法",
        "把 AP 想成一个会不断“看见、回忆、预测、解释、感觉、行动、再学习”的人工心智。它的重点不是一次答对，而是每一步为什么这样想、为什么这样看、为什么这样写都能被回放。",
        "EFF6FF",
        PALETTE["blue"],
    )
    add_callout(
        doc,
        "给技术员的读法",
        "重点跟踪 RuntimeTickEvent、StatePool、SSP、B/C/C*、DraftGrid、经验流、reward/punish、innate rule、action competition。所有 UI、TTS、画布、记忆包都应作为感受器、行动器或视图层接入同一闭环。",
        "F0FDF4",
        PALETTE["green"],
    )
    add_callout(
        doc,
        "给论文和审查的读法",
        "用积极证据口径表述：每个实验说明它可以证明什么；每个边界说明下一步要证明什么。基础能力证明期要避免答案表、关键词硬门、独立识别捷径、前端假 tick 和学生侧 LLM 代答。",
        "FFF7ED",
        PALETTE["orange"],
    )

    doc.add_heading("核心图景速览", level=2)
    table = doc.add_table(rows=1, cols=3)
    heads = ["图景", "一句话", "工程落点"]
    for idx, head in enumerate(heads):
        cell = table.cell(0, idx)
        cell.text = head
        set_cell_shading(cell, "EAF1F6")
        set_cell_border(cell)
        for run in cell.paragraphs[0].runs:
            set_run_font(run)
            run.font.bold = True
    rows = [
        ("完整闭环", "现实、预测、记忆、行动、奖惩在每个 tick 里循环。", "RuntimeTickEvent 与经验流"),
        ("状态池 / SSP", "状态池表示此刻亮着什么；SSP 表示刚才怎样发生。", "type projection 与 occurrence flow"),
        ("B/C/C*", "召回现在相似的历史，再预测未来、追溯原因。", "召回包与状态池回灌"),
        ("视觉重建", "视焦点附近高清，周边稀疏，多 tick 逐步补全。", "patch payload + clarity + source mask"),
        ("二维草稿", "输出是行动，不是一次性生成整句。", "DraftGrid + SELF_DRAFT_GRID"),
        ("统一经验流", "真相源是一条追加经验流，索引只加速。", "append-only event log"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
            set_cell_border(cells[idx])
            for para in cells[idx].paragraphs:
                for run in para.runs:
                    set_run_font(run)
                    run.font.size = Pt(9)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    hint = doc.add_paragraph("提示：若打开 Word 后目录未自动刷新，请右键目录选择“更新域/更新整个目录”。")
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hint.runs:
        set_run_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*rgb(PALETTE["muted"]))
    doc.add_page_break()


def add_inline_figure(doc: Document, caption: str, path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.45))
    p = doc.add_paragraph(caption, style="AP Figure Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def figure_insertion_anchors(figures: list[tuple[str, Path]]) -> list[tuple[str, str, Path]]:
    by_caption = {caption: path for caption, path in figures}
    return [
        ("1. AP 是什么", "图 1. AP 完整认知闭环", by_caption["AP 完整认知闭环"]),
        ("9. 状态池", "图 2. 状态池与短期结构池", by_caption["状态池与短期结构池"]),
        ("19. 统一经验流数据库", "图 6. 统一经验流与加速索引", by_caption["统一经验流与加速索引"]),
        ("56. B/C/C* 召回公式细化", "图 3. B/C/C* 召回与理解", by_caption["B/C/C* 召回与理解"]),
        ("第十六卷 视觉感受器", "图 4. 视焦点采样与内心画面重建", by_caption["视焦点采样与内心画面重建"]),
        ("第十八卷 DraftGrid", "图 5. DraftGrid 与二维草稿", by_caption["DraftGrid 与二维草稿"]),
        ("第二十卷 APV3 Phase20.7", "图 7. Phase20.7 实施路线", by_caption["Phase20.7 实施路线"]),
    ]


def sanitize_public_line(line: str) -> str:
    replacements = {
        "AP底层哲学、数学流程、实现红线、证据边界的当前最高标准草案": "AP底层哲学、数学流程、实现红线、证据边界的当前理论纠偏正本",
        "用户 2026-06-26 最新口径.": "最新理论裁定.",
        "第八卷 2026-06-26 最新理论裁定": "第八卷 最新理论裁定",
        "第二十六卷 Phase20.7 正式设计审查稿": "第二十六卷 Phase20.7 正式设计稿",
        "第二十卷 APV3 Phase20.7 设计草案骨架": "第二十卷 APV3 Phase20.7 设计骨架",
        "第三十卷 v0.4 Claude 独立审查吸收与哲学纠偏": "第三十卷 v0.4 最终理论纠偏与 AP-native 收束",
        "本卷吸收 Claude 独立对抗审查中成立的 10 项问题, 同时对其中部分修法做 AP-native 哲学纠偏.": "本卷收束 10 项高风险理论问题, 并把所有修订统一回 AP-native 哲学路线.",
        "本卷处理 v0.2 对抗性审查中提出的七个高风险点, 并按用户最新纠偏重新收束:": "本卷处理七个高风险理论点, 并按拟人性优先原则重新收束:",
        "186.3 对 Claude 可转发要点": "186.3 理论标准摘要",
        "让审查者能检查它是否真的走 AP-native 路线.": "让验证者能检查它是否真的走 AP-native 路线.",
        "审查约束": "验证约束",
        "设计审查": "设计校验",
        "对抗性审查": "理论验证",
        "对抗审查": "理论验证",
        "审查吸收": "理论吸收",
        "审查中": "校验中",
        "审查": "验证",
        "草案": "正本",
        "Claude": "理论验证",
        "Codex": "工程实现者",
        "GLM": "理论验证",
        "glm": "理论验证",
        "转发": "对外说明",
    }
    cleaned = line
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)

    q_names = [
        ("Q10", "问题十"),
        ("Q9", "问题九"),
        ("Q8", "问题八"),
        ("Q7", "问题七"),
        ("Q6", "问题六"),
        ("Q5", "问题五"),
        ("Q4", "问题四"),
        ("Q3", "问题三"),
        ("Q2", "问题二"),
        ("Q1", "问题一"),
    ]
    for old, new in q_names:
        cleaned = re.sub(rf"\b{old}\b", new, cleaned)
    cleaned = re.sub(r"问题([一二三四五六七八九十]+)\s+成立[,，]\s*但修法不是", r"问题\1的裁定：修法不是", cleaned)
    cleaned = re.sub(r"问题([一二三四五六七八九十]+)\s+成立[,，]\s*已补", r"问题\1的裁定：补入", cleaned)
    cleaned = re.sub(r"问题([一二三四五六七八九十]+)\s+成立[,，]\s*本卷补", r"问题\1的裁定：本卷补入", cleaned)
    cleaned = re.sub(r"问题([一二三四五六七八九十]+)\s+成立", r"问题\1的裁定", cleaned)
    cleaned = cleaned.replace("并吸收 理论验证 独立理论验证中成立的结构、公式、性能与实施映射修订", "并收束结构、公式、性能与实施映射修订")
    cleaned = cleaned.replace("理论验证 独立理论验证", "理论验证")
    cleaned = cleaned.replace("理论验证 可对外说明要点", "理论标准摘要")
    cleaned = cleaned.replace("对外理论校验要点", "理论标准摘要")
    cleaned = cleaned.replace("对外说明要点", "理论标准摘要")
    cleaned = cleaned.replace("理论校验", "理论验证")
    return cleaned


def maybe_add_anchored_figure(
    doc: Document,
    heading_text: str,
    anchors: list[tuple[str, str, Path]],
    inserted_figures: set[str],
) -> None:
    for anchor, caption, path in anchors:
        if anchor in heading_text and caption not in inserted_figures:
            add_inline_figure(doc, caption, path)
            inserted_figures.add(caption)
            return


def source_body_lines(lines: list[str]) -> list[str]:
    occurrences = [idx for idx, line in enumerate(lines) if line.strip() == "第一卷 总论与哲学基线"]
    if len(occurrences) >= 2:
        return lines[occurrences[1] :]
    if occurrences:
        return lines[occurrences[0] :]
    return lines


def is_heading2(lines: list[str], idx: int) -> bool:
    text = lines[idx].strip()
    if not H2_RE.match(text) or H3_RE.match(text):
        return False
    prev_blank = idx == 0 or not lines[idx - 1].strip()
    next_blank = idx + 1 >= len(lines) or not lines[idx + 1].strip()
    return prev_blank and next_blank and len(text) <= 80


def is_heading3(lines: list[str], idx: int) -> bool:
    text = lines[idx].strip()
    if not H3_RE.match(text):
        return False
    prev_blank = idx == 0 or not lines[idx - 1].strip()
    next_blank = idx + 1 >= len(lines) or not lines[idx + 1].strip()
    return prev_blank and next_blank and len(text) <= 90


def is_code_line(line: str) -> bool:
    if not line.strip():
        return False
    if line.startswith("  ") or line.startswith("\t"):
        return True
    stripped = line.strip()
    if len(stripped) <= 120 and CODE_HINT_RE.search(stripped) and not re.match(r"^[一-龥].*[。；，]$", stripped):
        return True
    return False


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F5F7")
    set_cell_border(cell, "D9E2EC")
    p = cell.paragraphs[0]
    p.style = "AP Code Block"
    p.paragraph_format.left_indent = Cm(0)
    text = "\n".join(line.rstrip() for line in code_lines)
    run = p.add_run(text)
    set_run_font(run, "Microsoft YaHei", "Consolas")
    run.font.size = Pt(8.6)


def add_body_paragraph(doc: Document, text: str) -> None:
    stripped = text.strip()
    if not stripped:
        return
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4)
    if re.match(r"^[0-9]+\.\s", stripped):
        p.paragraph_format.first_line_indent = Cm(0)
    else:
        p.paragraph_format.first_line_indent = Cm(0.72)
    run = p.add_run(stripped)
    set_run_font(run)
    run.font.size = Pt(10.5)


def add_full_whitepaper(doc: Document, lines: list[str], figures: list[tuple[str, Path]]) -> None:
    doc.add_heading("完整正文", level=1)
    note = doc.add_paragraph("以下为白皮书 30 卷完整正文，标题层级、公式块、长段落和核心图示已按 Word 阅读习惯重新排版。")
    set_paragraph_spacing(note)
    doc.add_page_break()

    code_buffer: list[str] = []
    first_volume = True
    figure_anchors = figure_insertion_anchors(figures)
    inserted_figures: set[str] = set()
    i = 0
    while i < len(lines):
        raw = sanitize_public_line(lines[i].rstrip("\n"))
        text = raw.strip()

        if not text:
            if code_buffer:
                add_code_block(doc, code_buffer)
                code_buffer = []
            i += 1
            continue

        if VOLUME_RE.match(text):
            if code_buffer:
                add_code_block(doc, code_buffer)
                code_buffer = []
            if not first_volume:
                doc.add_page_break()
            first_volume = False
            doc.add_heading(text, level=1)
            maybe_add_anchored_figure(doc, text, figure_anchors, inserted_figures)
            i += 1
            continue

        if is_heading3(lines, i):
            if code_buffer:
                add_code_block(doc, code_buffer)
                code_buffer = []
            doc.add_heading(text, level=3)
            maybe_add_anchored_figure(doc, text, figure_anchors, inserted_figures)
            i += 1
            continue

        if is_heading2(lines, i):
            if code_buffer:
                add_code_block(doc, code_buffer)
                code_buffer = []
            doc.add_heading(text, level=2)
            maybe_add_anchored_figure(doc, text, figure_anchors, inserted_figures)
            i += 1
            continue

        if is_code_line(raw):
            code_buffer.append(raw)
            i += 1
            continue

        if code_buffer:
            add_code_block(doc, code_buffer)
            code_buffer = []
        add_body_paragraph(doc, raw)
        i += 1

    if code_buffer:
        add_code_block(doc, code_buffer)


def add_back_matter(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("使用说明与维护建议", level=1)
    add_callout(
        doc,
        "如何作为论文底层纠偏文档使用",
        "优先引用本白皮书中的定义、闭环图、状态池/SSP 边界、B/C/C* 召回、视觉重建、DraftGrid、统一经验流和 Phase20.7 路线。论文表述尽量说明“该实验可以证明什么”，并把工程外壳与 AP-Core 证明边界分开。",
        "EFF6FF",
        PALETTE["blue"],
    )
    add_callout(
        doc,
        "如何作为工程实现标准使用",
        "新增功能前先判断它属于感受器、行动器、状态池视图、经验流索引、教师输入、审计旁路还是产品外壳。若无法接入同一 tick 闭环，应先回到白皮书重新设计。",
        "F0FDF4",
        PALETTE["green"],
    )
    add_callout(
        doc,
        "如何更新本 Word 版",
        "修改源 TXT 后运行 APV3.0test/scripts/docs/build_ap_whitepaper_docx.py 即可重建。图示为脚本生成，路径固定在 docs/assets/ap_whitepaper_figures。",
        "FFF7ED",
        PALETTE["orange"],
    )


def main() -> None:
    if not SOURCE_TXT.exists():
        raise FileNotFoundError(SOURCE_TXT)
    raw_lines = SOURCE_TXT.read_text(encoding="utf-8").splitlines()
    version_line = next((line for line in raw_lines[:20] if line.startswith("版本:")), "版本: 2026-06-26 冷保存正本 v0.4")
    figures = build_figures()

    doc = configure_document()
    add_cover(doc, version_line)
    add_reader_guide(doc)
    add_toc(doc)
    add_full_whitepaper(doc, source_body_lines(raw_lines), figures)
    add_back_matter(doc)
    request_update_fields(doc)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"saved={OUTPUT_DOCX}")
    print(f"figures={len(figures)} dir={FIGURE_DIR}")
    print(f"source_lines={len(raw_lines)}")


if __name__ == "__main__":
    main()
